import docx

doc = docx.Document("Word test 1.docx")
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[2].runs[1].underline = True
doc.paragraphs[4].style = 'Style_test'
doc.save("test.docx")